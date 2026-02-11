"""
Document loaders for different file formats.
"""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import List, Optional

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

from ..models import Document
from ..utils.logger import get_logger
from ..utils.text import clean_text

logger = get_logger(__name__)


class DocumentLoader(ABC):
    """Base class for document loaders."""
    
    @abstractmethod
    def load(self, file_path: str) -> Document:
        """Load a document from file."""
        pass
    
    @abstractmethod
    def supports(self, file_path: str) -> bool:
        """Check if loader supports this file type."""
        pass


class TextLoader(DocumentLoader):
    """Loader for plain text files."""
    
    def supports(self, file_path: str) -> bool:
        """Check if file is a text file."""
        return Path(file_path).suffix.lower() in ['.txt', '.text']
    
    def load(self, file_path: str) -> Document:
        """Load text file."""
        path = Path(file_path)
        
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        return Document(
            content=clean_text(content),
            source=str(path),
            metadata={
                'source': str(path),
                'file_type': 'text',
                'file_name': path.name,
            }
        )


class PDFLoader(DocumentLoader):
    """Loader for PDF files."""
    
    def supports(self, file_path: str) -> bool:
        """Check if file is a PDF."""
        return Path(file_path).suffix.lower() == '.pdf'
    
    def load(self, file_path: str) -> Document:
        """Load PDF file."""
        path = Path(file_path)
        
        try:
            reader = PdfReader(str(path))
            content_parts = []
            
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text:
                    content_parts.append(text)
            
            content = '\n\n'.join(content_parts)
            
            return Document(
                content=clean_text(content),
                source=str(path),
                metadata={
                    'source': str(path),
                    'file_type': 'pdf',
                    'file_name': path.name,
                    'num_pages': len(reader.pages),
                }
            )
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise


class DocxLoader(DocumentLoader):
    """Loader for DOCX files."""
    
    def supports(self, file_path: str) -> bool:
        """Check if file is a DOCX."""
        return Path(file_path).suffix.lower() in ['.docx', '.doc']
    
    def load(self, file_path: str) -> Document:
        """Load DOCX file."""
        path = Path(file_path)
        
        try:
            doc = DocxDocument(str(path))
            content_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        content_parts.append(row_text)
            
            content = '\n\n'.join(content_parts)
            
            return Document(
                content=clean_text(content),
                source=str(path),
                metadata={
                    'source': str(path),
                    'file_type': 'docx',
                    'file_name': path.name,
                    'num_paragraphs': len(doc.paragraphs),
                }
            )
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            raise


class MarkdownLoader(DocumentLoader):
    """Loader for Markdown files."""
    
    def supports(self, file_path: str) -> bool:
        """Check if file is Markdown."""
        return Path(file_path).suffix.lower() in ['.md', '.markdown']
    
    def load(self, file_path: str) -> Document:
        """Load Markdown file."""
        path = Path(file_path)
        
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        return Document(
            content=content,  # Keep markdown formatting
            source=str(path),
            metadata={
                'source': str(path),
                'file_type': 'markdown',
                'file_name': path.name,
            }
        )


class HTMLLoader(DocumentLoader):
    """Loader for HTML files."""
    
    def supports(self, file_path: str) -> bool:
        """Check if file is HTML."""
        return Path(file_path).suffix.lower() in ['.html', '.htm']
    
    def load(self, file_path: str) -> Document:
        """Load HTML file."""
        path = Path(file_path)
        
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()
        
        # Parse HTML and extract text
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(['script', 'style']):
            script.decompose()
        
        text = soup.get_text()
        
        return Document(
            content=clean_text(text),
            source=str(path),
            metadata={
                'source': str(path),
                'file_type': 'html',
                'file_name': path.name,
            }
        )


class CodeLoader(DocumentLoader):
    """Loader for source code files."""
    
    CODE_EXTENSIONS = [
        '.py', '.js', '.jsx', '.ts', '.tsx', '.java', '.cpp', '.c', '.h',
        '.cs', '.go', '.rs', '.rb', '.php', '.swift', '.kt', '.scala',
        '.r', '.sql', '.sh', '.bash', '.yaml', '.yml', '.json', '.xml'
    ]
    
    def supports(self, file_path: str) -> bool:
        """Check if file is a code file."""
        return Path(file_path).suffix.lower() in self.CODE_EXTENSIONS
    
    def load(self, file_path: str) -> Document:
        """Load code file."""
        path = Path(file_path)
        
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        return Document(
            content=content,  # Keep code formatting
            source=str(path),
            metadata={
                'source': str(path),
                'file_type': 'code',
                'file_name': path.name,
                'language': path.suffix.lstrip('.'),
            }
        )


class DocumentLoaderFactory:
    """Factory for creating appropriate document loaders."""
    
    def __init__(self):
        """Initialize with default loaders."""
        self.loaders: List[DocumentLoader] = [
            PDFLoader(),
            DocxLoader(),
            MarkdownLoader(),
            HTMLLoader(),
            CodeLoader(),
            TextLoader(),  # Fallback to text loader
        ]
    
    def add_loader(self, loader: DocumentLoader):
        """Add a custom loader."""
        self.loaders.insert(0, loader)  # Add to front for priority
    
    def load_file(self, file_path: str) -> Document:
        """Load a file using appropriate loader."""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not path.is_file():
            raise ValueError(f"Not a file: {file_path}")
        
        for loader in self.loaders:
            if loader.supports(file_path):
                logger.info(f"Loading {path.name} with {loader.__class__.__name__}")
                return loader.load(file_path)
        
        raise ValueError(f"No loader found for file: {file_path}")
    
    def load_directory(
        self,
        directory: str,
        recursive: bool = True,
        file_pattern: Optional[str] = None,
    ) -> List[Document]:
        """Load all documents from a directory."""
        dir_path = Path(directory)
        
        if not dir_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        if not dir_path.is_dir():
            raise ValueError(f"Not a directory: {directory}")
        
        documents = []
        pattern = file_pattern or '*'
        
        if recursive:
            files = dir_path.rglob(pattern)
        else:
            files = dir_path.glob(pattern)
        
        for file_path in files:
            if file_path.is_file():
                try:
                    doc = self.load_file(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    logger.warning(f"Failed to load {file_path}: {e}")
        
        logger.info(f"Loaded {len(documents)} documents from {directory}")
        return documents
